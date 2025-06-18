from flask import Flask, render_template, request, make_response
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import cohere

app = Flask(__name__)

# Set your Cohere API key
co = cohere.Client("NqCDyPmfZHiXEDiyn0Xooutz67b0XHFPoeZ8qeYy")  # Replace with your real key

generated_data = {}

@app.route("/")
def home():
    return render_template("index.html")

@app.route("/generate", methods=["POST"])
def generate():
    global generated_data
    data = request.form.to_dict()

    prompt = (
        f"Create a clean, professional resume for {data.get('name')} applying as a {data.get('job')}.\n"
        f"Email: {data.get('email')}\n"
        f"Phone: {data.get('phone')}\n"
        f"Summary: {data.get('summary')}\n"
        f"Skills: {data.get('skills')}\n"
        f"Work Experience: {data.get('experience')}\n"
        f"Education: {data.get('education')}\n"
        f"Certifications: {data.get('certifications', '')}\n"
        "Format with clear headings and bullet points. Do NOT include references."
    )

    response = co.generate(
        model="command",
        prompt=prompt,
        max_tokens=1500,
        temperature=0.7,
    )

    resume_text = response.generations[0].text.strip()

    generated_data = {
        "name": data.get("name", "Your Name"),
        "job": data.get("job", ""),
        "email": data.get("email", ""),
        "phone": data.get("phone", ""),
        "resume_text": resume_text,
    }

    return render_template("result.html", resume=resume_text, **generated_data)

@app.route("/download-word")
def download_word():
    global generated_data
    doc = Document()

    # Title
    title = doc.add_heading(generated_data.get("name", "Your Name"), level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Job title
    job_title = doc.add_paragraph(generated_data.get("job", ""))
    job_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    job_title.runs[0].font.size = Pt(14)
    job_title.runs[0].bold = True

    # Contact
    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = contact.add_run(f"Email: {generated_data.get('email', '')} | Phone: {generated_data.get('phone', '')}")
    run.font.size = Pt(10)
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    def add_heading(text):
        heading = doc.add_heading(text.upper(), level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_bullets(text):
        lines = [line.strip("•- \t") for line in text.split("\n") if line.strip()]
        for line in lines:
            doc.add_paragraph(line, style="List Bullet")

    known_sections = [
        "SUMMARY", "SKILLS", "WORK EXPERIENCE", "ENGINEERING PROJECTS",
        "RELEVANT SKILLS", "ABILITIES", "EDUCATION", "CERTIFICATIONS", "CONTACT INFORMATION"
    ]

    text = generated_data.get("resume_text", "")
    sections = {}
    current_section = None

    for line in text.splitlines():
        stripped = line.strip()
        heading_candidate = stripped.rstrip(":").upper()
        if heading_candidate in known_sections:
            current_section = heading_candidate
            sections[current_section] = ""
        elif current_section:
            sections[current_section] += line + "\n"

    for section in known_sections:
        content = sections.get(section, "").strip()
        if content:
            add_heading(section)
            if section in ["SUMMARY", "EDUCATION", "CONTACT INFORMATION"]:
                doc.add_paragraph(content)
            else:
                add_bullets(content)
            doc.add_paragraph()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = make_response(buffer.read())
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response.headers["Content-Disposition"] = "attachment; filename=resume.docx"
    return response

if __name__ == "__main__":
    app.run(debug=True)
